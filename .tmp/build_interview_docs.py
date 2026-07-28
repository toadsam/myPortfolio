# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "interview-docs"


def repo_snapshot(repo: str, commit: str, date: str, message: str) -> str:
    return f"{repo} / {commit} / {date} / {message}"


PROJECTS = [
    {
        "filename": "MuscleUp_면접대비_기술문서.docx",
        "title": "득근득근 MuscleUp",
        "subtitle": "피트니스 커뮤니티, 운동 기록, 캐릭터 성장, 실시간 라운지를 결합한 풀스택 서비스",
        "snapshot": repo_snapshot(
            "https://github.com/toadsam/Ajou_MuscleUp",
            "38c2c2b",
            "2026-05-07",
            "Merge pull request #4 from toadsam/26.05.06",
        ),
        "one_liner": (
            "MuscleUp은 운동 기록을 단순 저장이 아니라 출석, 포인트, 캐릭터 성장, 랭킹, 커뮤니티, "
            "실시간 라운지로 이어지게 만든 피트니스 동기부여 플랫폼입니다."
        ),
        "problem": [
            "운동 앱은 기록을 남기더라도 지속 동기가 약해 이탈이 빠르게 발생합니다.",
            "사용자는 오늘의 운동, 누적 출석, 성장 상태, 친구/크루 활동을 한 화면에서 보고 싶어 합니다.",
            "커뮤니티와 랭킹은 참여 동기를 주지만, 인증/권한/실시간 상태 관리가 함께 필요합니다.",
            "InBody 이미지/PDF 기반 AI 피드백은 파일 처리, 품질 검사, 응답 구조화, 보고서 생성이 같이 얽힙니다.",
            "프론트엔드, Spring Boot API, Socket.IO 실시간 서버가 서로 다른 관심사를 나누도록 설계했습니다.",
        ],
        "architecture": [
            "React/Vite 프론트엔드는 라우팅, 인증 상태, API 호출, 차트, PDF 다운로드, PWA 경험을 담당합니다.",
            "Spring Boot 백엔드는 인증, 출석, 캐릭터, 크루, AI, 파일, 이벤트, 관리자 기능을 REST API로 제공합니다.",
            "Node Socket.IO 실시간 서버는 라운지 접속자, 위치 이동, 채팅, 이모트 브로드캐스트를 담당합니다.",
            "JWT access token과 refresh token을 쿠키/Authorization 양쪽에서 처리하여 브라우저 제약을 완화합니다.",
            "운동 기록은 출석 점수, 캐릭터 성장, 공유 링크, 주간 랭킹으로 전파됩니다.",
        ],
        "diagram": """
        [Browser React]
            | Axios/fetch + accessToken + cookies
            v
        [Spring Boot API] ---- JPA ---- [DB: users, attendance, crew, character, refresh_tokens]
            | files / AI report
            v
        [OpenAI + file processing + PDFBox]

        [Browser Socket.IO] <----> [Node realtime server]
            | join / move / chat / emote
            v
        [in-memory rooms + throttled broadcasts]
        """,
        "dirs": [
            ("backend", "Spring Boot 3.5 API. 인증, 출석, AI, 크루, 캐릭터, 파일, 관리자 기능이 들어 있습니다."),
            ("frontend", "React 19, TypeScript, Vite 기반 클라이언트. 라우팅, TanStack Query, Axios, 차트, PWA를 사용합니다."),
            ("realtime", "Socket.IO + TypeScript 실시간 라운지 서버. 위치/채팅/타이핑/이모트 이벤트를 처리합니다."),
            ("docs", "프로젝트 문서와 보조 자료가 들어 있습니다."),
            ("uploads", "개발 환경 파일 업로드/임시 리소스가 위치합니다."),
        ],
        "stack": [
            ("Frontend", "React 19, TypeScript, Vite, React Router 7, TanStack Query, Axios, Recharts, Socket.IO client"),
            ("Backend", "Java 17, Spring Boot 3.5, Spring Security, JPA, Validation, JavaMail, Google API, PDFBox"),
            ("Realtime", "Node.js, Socket.IO, TypeScript, tsx"),
            ("Data/Infra", "MySQL/PostgreSQL, AWS S3 SDK, JWT, refresh token table"),
            ("AI/File", "OpenAI chat completion, image/PDF preprocessing, PDF report generation"),
        ],
        "domain": [
            ("User", "이메일/Google OAuth 로그인 주체이며 JWT subject로 사용됩니다."),
            ("RefreshToken", "장기 세션 유지와 access token 재발급을 담당합니다."),
            ("Attendance", "운동 기록의 원장 역할을 하며 성장/랭킹/공유로 이어집니다."),
            ("Character", "운동 누적 포인트와 신체 정보에 따라 성장 상태가 변합니다."),
            ("Crew", "초대 코드, 가입 요청, 챌린지, 출석 기반 점수를 가진 그룹 단위입니다."),
            ("Realtime player", "Socket 연결 단위의 닉네임, 레벨, 티어, 위치, 상태 정보입니다."),
        ],
        "flows": [
            {
                "title": "로그인과 토큰 재발급 흐름",
                "steps": [
                    "사용자가 이메일/비밀번호 또는 Google ID token으로 로그인합니다.",
                    "백엔드는 access token과 refresh token을 발급하고, refresh/access/rememberMe 쿠키를 설정합니다.",
                    "프론트 Axios 인터셉터는 localStorage access token을 Authorization 헤더에 붙입니다.",
                    "API가 401을 반환하면 단일 refresh 요청만 보내고, 대기 중인 요청은 큐에 넣었다가 새 token으로 재시도합니다.",
                    "로그아웃 시 refresh token을 폐기하고 관련 쿠키를 제거합니다.",
                ],
                "interview": "인증을 설명할 때는 'JWT를 썼다'에서 멈추지 말고, 401 동시 발생 시 refresh stampede를 막은 점을 말하면 좋습니다.",
            },
            {
                "title": "운동 기록에서 캐릭터 성장까지",
                "steps": [
                    "사용자가 오늘의 운동 강도와 기록을 저장합니다.",
                    "AttendanceService가 한국 시간 기준으로 오늘 기록을 upsert하고, 허용된 강도/타입인지 검증합니다.",
                    "운동 강도별 포인트를 계산하고 이벤트 진행도, 공유 slug, 응원/리포트 정보를 업데이트합니다.",
                    "CharacterGrowthService가 누적 포인트와 UserBodyStats를 기반으로 CharacterService.evaluate를 호출합니다.",
                    "결과적으로 기록 저장이 단순 CRUD가 아니라 성장, 랭킹, 공유 기능으로 연결됩니다.",
                ],
                "interview": "이 흐름은 도메인 이벤트를 별도 메시지 큐로 분리하지 않았지만, 서비스 간 책임을 나눠 확장 가능하게 만든 예시입니다.",
            },
            {
                "title": "실시간 라운지 이벤트 흐름",
                "steps": [
                    "클라이언트가 Socket.IO로 접속하고 join payload를 보냅니다.",
                    "서버는 닉네임, 레벨, 티어, 스테이지, 선택 필드를 검증한 뒤 방에 플레이어를 등록합니다.",
                    "move/chat/typing/emote/sticker 이벤트를 받아 해당 방 사용자에게 브로드캐스트합니다.",
                    "위치 이동은 60ms 주기로 묶어서 보내 과도한 브로드캐스트를 줄입니다.",
                    "현재 구현에는 JWT 검증 TODO가 있어 운영 전 보안 강화를 설명할 수 있습니다.",
                ],
                "interview": "실시간 서버는 상태를 DB에 매번 쓰지 않고 in-memory room으로 관리해 지연을 줄인 설계입니다.",
            },
        ],
        "code": [
            {
                "file": "backend/src/main/java/com/ajou/muscleup/config/SecurityConfig.java",
                "snippet": """
                http
                    .csrf(csrf -> csrf.disable())
                    .sessionManagement(sm -> sm.sessionCreationPolicy(SessionCreationPolicy.STATELESS))
                    .authorizeHttpRequests(auth -> auth
                        .requestMatchers("/api/auth/**", "/api/email/**", "/api/events/**").permitAll()
                        .requestMatchers("/api/admin/**").hasRole("ADMIN")
                        .anyRequest().authenticated())
                    .addFilterBefore(jwtAuthenticationFilter, UsernamePasswordAuthenticationFilter.class)
                    .addFilterAfter(auditLogFilter, JwtAuthenticationFilter.class);
                """,
                "explain": [
                    "세션을 만들지 않는 stateless API로 구성해 JWT 기반 인증 모델과 맞췄습니다.",
                    "공개 엔드포인트, 일반 사용자, 관리자 엔드포인트를 SecurityFilterChain에서 분리합니다.",
                    "JWT 필터를 UsernamePasswordAuthenticationFilter 앞에 두어 컨트롤러 진입 전 SecurityContext를 구성합니다.",
                    "감사 로그 필터를 뒤에 붙여 인증된 요청의 흔적을 남길 수 있게 했습니다.",
                ],
            },
            {
                "file": "frontend/src/lib/api.ts",
                "snippet": """
                let refreshing = false;
                const waiters: Array<(token: string | null) => void> = [];

                api.interceptors.response.use(undefined, async (error) => {
                  if (error.response?.status !== 401 || error.config._retry) throw error;
                  error.config._retry = true;
                  if (refreshing) {
                    const token = await new Promise(resolve => waiters.push(resolve));
                    error.config.headers.Authorization = `Bearer ${token}`;
                    return api(error.config);
                  }
                  refreshing = true;
                  const refreshed = await api.post("/api/auth/refresh");
                  waiters.splice(0).forEach(resolve => resolve(refreshed.data.accessToken));
                  return api(error.config);
                });
                """,
                "explain": [
                    "여러 요청이 동시에 401을 받으면 refresh API를 여러 번 때리는 문제가 생깁니다.",
                    "refreshing 플래그와 waiter 큐를 두어 첫 요청만 refresh를 수행하고 나머지는 결과를 기다립니다.",
                    "새 access token을 받은 뒤 원래 요청을 재시도해 사용자는 세션 만료를 덜 체감합니다.",
                    "면접에서는 '동시성 이슈를 프론트 레이어에서도 고려했다'는 점을 강조할 수 있습니다.",
                ],
            },
            {
                "file": "realtime/src/server.ts",
                "snippet": """
                socket.on("move", (payload) => {
                  const player = updatePlayer(socket.id, {
                    x: clamp(payload.x, 0, MAP_WIDTH),
                    y: clamp(payload.y, 0, MAP_HEIGHT),
                    direction: payload.direction,
                  });
                  pendingBroadcast.add(player.roomId);
                });

                setInterval(() => {
                  for (const roomId of pendingBroadcast) {
                    io.to(roomId).emit("players", listPlayers(roomId));
                  }
                  pendingBroadcast.clear();
                }, 60);
                """,
                "explain": [
                    "클라이언트 이동 이벤트를 받을 때마다 즉시 전체 상태를 보내면 네트워크 부하가 커집니다.",
                    "방 ID를 pendingBroadcast에 모아두고 60ms마다 한 번만 최신 상태를 전송합니다.",
                    "좌표는 서버에서 clamp해 비정상 좌표가 맵 밖으로 나가지 않게 합니다.",
                    "MVP에서는 in-memory 상태로 충분하지만, 다중 인스턴스 운영 시 Redis adapter가 다음 개선점입니다.",
                ],
            },
        ],
        "frontend": [
            "ProtectedRoute는 `/api/auth/me`를 확인하지만 네트워크/Safari 계열 문제에서 즉시 로그아웃시키지 않는 완충 로직을 둡니다.",
            "Axios뿐 아니라 native fetch를 사용하는 코드까지 `installFetchAuth`로 감싸 토큰 재발급 정책을 맞춥니다.",
            "운동, 랭킹, AI 리포트 화면은 React Query와 차트/다운로드 라이브러리를 함께 사용합니다.",
            "Socket.IO 클라이언트는 HTTP API와 별도 채널로 실시간 라운지 UX를 만듭니다.",
        ],
        "security": [
            "JWT secret 길이를 32자 이상 요구해 HMAC key 안정성을 확보합니다.",
            "access token은 Authorization 헤더와 cookie 양쪽에서 읽어 배포 환경별 브라우저 제약을 완화합니다.",
            "refresh token service는 rotate라는 이름이지만 병렬 refresh 로그아웃을 막기 위해 같은 refresh token을 유지하는 안정화 전략입니다.",
            "실시간 서버의 JWT 검증 TODO는 운영 전 반드시 보완해야 할 리스크입니다.",
        ],
        "ops": [
            "AI 파일 처리에는 이미지/PDF 전처리, 용량 검증, 품질 점수, JSON 구조화, PDFBox 보고서 생성이 포함됩니다.",
            "파일 저장소는 로컬/클라우드 환경 차이를 고려해야 하며 S3 SDK 의존성이 포함되어 있습니다.",
            "실시간 서버는 Socket.IO health/status endpoint로 별도 모니터링할 수 있습니다.",
        ],
        "tests": [
            "인증은 access 만료, refresh 성공, refresh 실패, 로그아웃 후 재요청 케이스를 테스트해야 합니다.",
            "출석은 한국 시간 기준 날짜 경계, 같은 날 upsert, 강도별 포인트 계산을 검증해야 합니다.",
            "실시간 라운지는 join payload 검증, 좌표 clamp, 방별 브로드캐스트 분리를 테스트하는 것이 좋습니다.",
        ],
        "strengths": [
            "운동 기록을 성장/랭킹/커뮤니티로 연결해 기능 간 시너지가 분명합니다.",
            "프론트와 백엔드 모두 토큰 만료 처리를 신경 써 실제 사용자 경험이 안정적입니다.",
            "AI, 파일, 실시간, 커뮤니티 등 여러 기술 요소를 하나의 서비스 목표에 묶었습니다.",
        ],
        "improvements": [
            "Socket.IO 서버에서 JWT를 검증하고 사용자 ID 기반으로 방 권한을 확인해야 합니다.",
            "refresh token 정책은 이름과 구현을 맞추거나, 안정 refresh 전략이라고 명시하는 편이 좋습니다.",
            "실시간 서버 다중 인스턴스 운영을 위해 Redis adapter와 sticky session 전략을 준비할 수 있습니다.",
            "AI 결과의 schema validation과 retry 정책을 테스트 코드로 고정하면 회귀를 줄일 수 있습니다.",
        ],
        "qa": [
            ("왜 Spring Boot와 Socket.IO 서버를 나눴나요?", "REST 도메인 API와 실시간 위치/채팅 처리는 부하 특성과 상태 모델이 달라 분리했습니다. Spring은 인증/DB 트랜잭션에 집중하고, Node는 낮은 지연의 이벤트 브로드캐스트에 집중합니다."),
            ("JWT refresh 동시성 문제는 어떻게 처리했나요?", "프론트에서 refresh 요청을 하나로 묶고 나머지 요청은 큐에 대기시킨 뒤 새 access token으로 재시도합니다. 백엔드도 병렬 refresh 때문에 로그아웃되는 문제를 줄이도록 안정적으로 처리합니다."),
            ("가장 어려웠던 부분은?", "운동 기록이 여러 기능으로 퍼지는 부분입니다. 출석 저장이 캐릭터 성장, 공유, 랭킹에 영향을 주기 때문에 서비스 책임을 나누고 순서를 명확히 잡는 것이 중요했습니다."),
            ("운영 전에 반드시 고칠 점은?", "실시간 서버의 JWT 검증, refresh token 정책 명확화, AI 파일 처리 실패 케이스 테스트를 우선 보완하겠습니다."),
        ],
    },
    {
        "filename": "FestFlow_면접대비_기술문서.docx",
        "title": "FestFlow",
        "subtitle": "대학교 축제 부스, 혼잡도, 운영자, 예약, 공지, AI 예측을 다루는 축제 운영 웹앱",
        "snapshot": repo_snapshot(
            "https://github.com/toadsam/FestFlow",
            "c5ed07f",
            "2026-06-13",
            "[ADD] JAVA로 파이썬 모델 실행 가능하도록 -재훈",
        ),
        "one_liner": (
            "FestFlow는 축제 현장의 방문객 경험과 운영자 의사결정을 동시에 지원하기 위해 지도, 부스, GPS 로그, "
            "예약, SSE 실시간 스트림, AI 혼잡도 예측을 결합한 운영형 서비스입니다."
        ),
        "problem": [
            "축제 현장에서는 어느 부스가 혼잡한지, 어디로 이동해야 하는지 사용자가 즉시 알기 어렵습니다.",
            "운영자는 공지, 분실물, 예약, 스태프 상태, 부스 재고와 혼잡도를 빠르게 확인해야 합니다.",
            "GPS 로그와 예약 데이터는 실시간성이 중요하지만 WebSocket까지 필요하지 않은 단방향 알림도 많습니다.",
            "AI 혼잡도 예측은 Python 모델과 Java 서비스 사이의 실행/실패 처리 경계가 중요합니다.",
            "지도 기반 UI와 관리자/부스 운영자 UI가 서로 다른 권한과 사용 맥락을 가집니다.",
        ],
        "architecture": [
            "React/Vite/Tailwind 프론트는 일반 방문객 지도, 관리자 대시보드, 부스 운영자 페이지를 나눕니다.",
            "Spring Boot 백엔드는 부스/공지/예약/운영자 인증/혼잡도/AI 분석 API를 제공합니다.",
            "SSE는 혼잡도, 이벤트, 공지, 부스, 스태프, 분실물, 예약처럼 서버에서 클라이언트로 흐르는 갱신에 사용됩니다.",
            "GPS 로그와 예약 정보를 혼잡도 계산에 반영하고, Python 모델이 없거나 실패하면 portable fallback 모델을 사용합니다.",
            "운영자 API는 X-OPS-KEY와 JWT/role 기반 접근 제어가 같이 들어갑니다.",
        ],
        "diagram": """
        [Visitor Map]
           | GPS log / booth query / SSE subscribe
           v
        [Spring Boot API] ---- JPA ---- [Booths, GPS logs, Reservations, Staff, Notices]
           |                         |
           | SSE publish             | snapshot
           v                         v
        [Browser EventSource]   [AI Congestion Service]
                                  | run python model
                                  v
                              [portable fallback model]
        """,
        "dirs": [
            ("backend", "Spring Boot API. SSE, 혼잡도, 운영자 인증, 예약, 부스, AI 모델 실행이 포함됩니다."),
            ("frontend", "React 18, Vite, Tailwind, React Leaflet 기반 방문객/관리자/부스 운영자 UI입니다."),
            ("docs", "서비스 설명과 설계 문서가 들어 있습니다."),
            ("exports", "데이터/모델 출력물 또는 운영 보조 파일을 보관하는 폴더입니다."),
            ("scripts/tools", "모델 실행과 보조 자동화 스크립트가 위치합니다."),
        ],
        "stack": [
            ("Frontend", "React 18, Vite, Tailwind, React Router 6, React Leaflet, QRCode, EventSource"),
            ("Backend", "Java 17, Spring Boot 3.3, JPA, Security, Validation, JWT"),
            ("Realtime", "Server-Sent Events. 단방향 실시간 알림에 적합합니다."),
            ("AI", "Java service에서 Python congestion model을 실행하고 JSON 입출력으로 결과를 받습니다."),
            ("Data", "MySQL/PostgreSQL, GPS logs, reservations, booth snapshots, decision logs"),
        ],
        "domain": [
            ("Booth", "축제 지도에 표시되는 운영 단위이며 좌표, 메뉴, 재고, 예약 테이블을 가집니다."),
            ("GpsLog", "방문객 위치 로그. 최근 15분 로그를 기준으로 부스 반경 내 혼잡도를 계산합니다."),
            ("Reservation", "부스 테이블 예약과 체크인 상태. 좌석 압박 계산에 사용됩니다."),
            ("Staff", "현장 스태프 상태와 공지 대상입니다."),
            ("Stream", "브라우저 EventSource가 구독하는 SSE emitter 묶음입니다."),
            ("AI Snapshot", "현재 crowd/wait/reservation/stock/event 정보를 모델 입력으로 정규화한 데이터입니다."),
        ],
        "flows": [
            {
                "title": "GPS 로그에서 혼잡도 갱신까지",
                "steps": [
                    "방문객 지도 화면이 주기적으로 또는 이동 시 GPS 로그를 서버로 보냅니다.",
                    "GpsService는 로그를 저장한 뒤 BoothService.getAllCongestions를 호출합니다.",
                    "BoothService는 부스 반경 80m 이내의 최근 15분 GPS 로그를 필터링합니다.",
                    "최근 로그일수록 가중치를 높게 주고 weighted count를 계산합니다.",
                    "결과를 여유/보통/혼잡/매우혼잡으로 변환하고 SSE로 구독자에게 보냅니다.",
                ],
                "interview": "위치 데이터는 정확한 인원 수가 아니라 '현장 의사결정에 충분한 지표'로 설계했다는 점을 설명하면 좋습니다.",
            },
            {
                "title": "SSE 기반 실시간 업데이트",
                "steps": [
                    "프론트는 혼잡도, 공지, 부스, 예약 등 필요한 채널별 EventSource를 엽니다.",
                    "StreamService는 채널별 SseEmitter 리스트를 CopyOnWriteArrayList로 관리합니다.",
                    "업데이트가 발생하면 named event로 데이터를 전송합니다.",
                    "완료, 타임아웃, 에러 콜백에서 emitter를 제거해 누수를 줄입니다.",
                    "죽은 연결로 send가 실패하면 해당 emitter를 리스트에서 제거합니다.",
                ],
                "interview": "양방향 채팅이 아니라 서버 갱신을 화면에 밀어주는 구조라 SSE가 WebSocket보다 단순하고 적합했습니다.",
            },
            {
                "title": "AI 혼잡도 예측과 fallback",
                "steps": [
                    "현재 부스 상태, GPS 변화량, 예약 변화량, 재고, 이벤트 임박 여부를 snapshot으로 만듭니다.",
                    "Java 서비스가 PythonCongestionModelService를 통해 Python 스크립트를 실행합니다.",
                    "입력/출력은 임시 JSON 파일로 전달하며 timeout을 둡니다.",
                    "모델 파일이나 Python 실행이 실패하면 portable model JSON 또는 규칙 기반 fallback을 사용합니다.",
                    "운영 화면은 모델이 불안정해도 최소한의 위험도/권고를 계속 받을 수 있습니다.",
                ],
                "interview": "AI 기능을 붙일 때 모델 실패가 서비스 장애로 번지지 않게 fallback 경로를 설계한 점이 핵심입니다.",
            },
        ],
        "code": [
            {
                "file": "backend/src/main/java/com/festflow/backend/service/stream/StreamService.java",
                "snippet": """
                private final List<SseEmitter> congestionEmitters = new CopyOnWriteArrayList<>();

                public SseEmitter subscribeCongestion() {
                    SseEmitter emitter = new SseEmitter(0L);
                    congestionEmitters.add(emitter);
                    emitter.onCompletion(() -> congestionEmitters.remove(emitter));
                    emitter.onTimeout(() -> congestionEmitters.remove(emitter));
                    emitter.onError(error -> congestionEmitters.remove(emitter));
                    return emitter;
                }

                public void publishCongestion(Object payload) {
                    send(congestionEmitters, "congestion", payload);
                }
                """,
                "explain": [
                    "SSE 연결은 오래 유지되므로 리스트에서 emitter 생명주기를 관리해야 합니다.",
                    "CopyOnWriteArrayList는 연결 수가 매우 크지 않고 순회 중 제거가 필요한 상황에서 구현이 단순합니다.",
                    "채널을 나눠 두면 프론트가 필요한 이벤트만 구독할 수 있습니다.",
                    "면접에서는 '실시간이지만 양방향성이 필요하지 않아 SSE를 선택했다'고 말할 수 있습니다.",
                ],
            },
            {
                "file": "backend/src/main/java/com/festflow/backend/service/BoothService.java",
                "snippet": """
                private static final double BOOTH_RADIUS_METERS = 80.0;

                long weightedCount = recentLogs.stream()
                    .filter(log -> distanceInMeters(booth, log) <= BOOTH_RADIUS_METERS)
                    .mapToDouble(this::timeWeight)
                    .sum();

                if (count < 3) return "여유";
                if (count < 7) return "보통";
                if (count < 12) return "혼잡";
                return "매우혼잡";
                """,
                "explain": [
                    "혼잡도는 최근 15분 로그 중 부스 반경 안에 있는 로그를 기준으로 계산합니다.",
                    "시간 가중치를 둬 방금 들어온 로그가 오래된 로그보다 큰 영향을 주게 했습니다.",
                    "Haversine 거리 계산으로 위도/경도 좌표의 실제 거리를 근사합니다.",
                    "절대 정확도보다 현장 화면에서 빠르게 이해 가능한 4단계 상태가 목적입니다.",
                ],
            },
            {
                "file": "backend/src/main/java/com/festflow/backend/service/ai/PythonCongestionModelService.java",
                "snippet": """
                Process process = new ProcessBuilder(
                    pythonCommand, scriptPath, modelPath, inputJson, outputJson
                ).start();

                boolean finished = process.waitFor(timeoutSeconds, TimeUnit.SECONDS);
                if (!finished || process.exitValue() != 0) {
                    return portableFallback(input);
                }
                return readPrediction(outputJson);
                """,
                "explain": [
                    "Java 서비스에서 Python 모델을 직접 호출하되 timeout과 exit code를 확인합니다.",
                    "모델 실행 실패를 API 실패로 그대로 노출하지 않고 fallback 예측으로 전환합니다.",
                    "입출력을 JSON 파일로 분리해 언어 간 경계를 명확히 했습니다.",
                    "운영 환경에서는 모델 drift status와 로그를 함께 봐야 합니다.",
                ],
            },
        ],
        "frontend": [
            "StageMapPage는 React Leaflet 지도에서 부스 마커, 검색/필터, 혼잡도 배지, GPS 전송을 처리합니다.",
            "API 모듈은 fetchWithTimeout과 EventSource 생성 함수를 분리해 네트워크 지연과 실시간 구독을 관리합니다.",
            "AdminPage는 booths/events/notices/KPI/logs/staff/AI 데이터를 Promise.allSettled로 불러와 일부 실패에도 화면을 구성합니다.",
            "OpsBoothPage는 부스 운영자가 키 기반 인증 후 예약 스트림, 메뉴, 이미지, 테이블 상태를 관리합니다.",
        ],
        "security": [
            "Spring Security는 OPS_MASTER, OPS_BOOTH, ADMIN 권한을 분리합니다.",
            "운영자 API는 `X-OPS-KEY` 필터로 boothId와 역할을 인증 컨텍스트에 넣습니다.",
            "일반 `/api/**` 중 공개가 많은 구조라 운영 전 민감 API permitAll 여부를 다시 점검해야 합니다.",
            "JWT filter와 OpsKey filter가 함께 있어 요청 경로별 인증 방식이 다릅니다.",
        ],
        "ops": [
            "SSE emitter는 연결 종료/타임아웃/에러 시 제거하지 않으면 메모리 누수와 dead connection 문제가 생깁니다.",
            "Python 모델 실행은 timeout, 모델 파일 존재 여부, portable fallback, drift status까지 운영 관점으로 설계되어야 합니다.",
            "지도 좌표는 캠퍼스 중심 반경 검증과 fallback 좌표로 데이터 오류에 대응합니다.",
        ],
        "tests": [
            "혼잡도 계산은 반경 80m 경계, 시간 가중치, 단계 변환 임계값을 단위 테스트로 고정해야 합니다.",
            "SSE는 emitter 제거와 send 실패 제거 경로를 테스트해야 합니다.",
            "Python 모델 실행은 성공, timeout, 파일 없음, 잘못된 JSON, fallback 케이스를 모두 검증해야 합니다.",
        ],
        "strengths": [
            "방문객 지도와 운영자 대시보드를 하나의 데이터 흐름으로 연결했습니다.",
            "SSE를 적절히 선택해 구현 복잡도를 낮추면서 실시간성을 확보했습니다.",
            "AI 모델 실행 실패를 고려한 fallback 구조가 실서비스 관점에서 강점입니다.",
        ],
        "improvements": [
            "공개 API와 권한 API의 permitAll 범위를 더 엄격히 분리해야 합니다.",
            "GPS 로그 기반 혼잡도는 중복 사용자/스푸핑 방어가 필요합니다.",
            "SSE 연결 수가 커지면 broker 또는 topic 기반 event bus 도입을 검토할 수 있습니다.",
            "AI 모델 입출력 스키마를 JSON Schema로 검증하면 장애 원인을 더 빨리 찾을 수 있습니다.",
        ],
        "qa": [
            ("왜 WebSocket 대신 SSE인가요?", "혼잡도와 공지처럼 서버가 클라이언트로 밀어주는 단방향 업데이트가 중심이라 SSE가 더 단순하고 HTTP 친화적입니다."),
            ("혼잡도는 어떻게 계산하나요?", "최근 15분 GPS 로그 중 부스 반경 80m 이내 로그를 골라 시간 가중치를 더하고, count 구간에 따라 4단계 상태로 바꿉니다."),
            ("AI 모델이 실패하면 어떻게 되나요?", "Python 실행 timeout, 모델 파일 부재, 출력 실패를 감지하고 portable fallback 또는 규칙 기반 계산으로 전환합니다."),
            ("가장 중요한 개선점은?", "민감 API의 권한 범위를 정리하고 GPS 신뢰성/악용 방어를 강화하는 것입니다."),
        ],
    },
    {
        "filename": "MyStockDesk_면접대비_기술문서.docx",
        "title": "MyStock-Desk / StockFlow",
        "subtitle": "투자 기록을 기반으로 포트폴리오, 리스크, 리서치, AI 체크포인트를 제공하는 주식 학습 데스크",
        "snapshot": repo_snapshot(
            "https://github.com/toadsam/MyStock-Desk",
            "6cb3d38",
            "2026-06-22",
            "[ADD] 분석 기능 추가 -재훈",
        ),
        "one_liner": (
            "StockFlow는 실제 주문/증권사 연동이 아니라 사용자가 직접 입력한 투자 기록을 원장으로 삼아 "
            "보유 비중, 수익률, 리스크 알림, 학습 후보, 뉴스/공시 영향, 체크리스트를 제공하는 투자 복기 도구입니다."
        ),
        "problem": [
            "초보 투자자는 매수/매도 이유를 기록하지 않아 나중에 의사결정을 복기하기 어렵습니다.",
            "포트폴리오 수익률은 현재가와 거래 원장을 함께 계산해야 하므로 단순 화면 합산으로는 부족합니다.",
            "외부 시세 API는 실패하거나 제한될 수 있어 데모/seed 데이터 fallback이 필요합니다.",
            "AI 성격의 기능은 매수/매도 추천이 아니라 위험 점검과 학습 보조로 경계를 명확히 해야 합니다.",
            "프론트는 실 API와 mock fallback 양쪽에서 화면이 무너지지 않아야 합니다.",
        ],
        "architecture": [
            "React/Vite/Tailwind 프론트는 포트폴리오, 거래 기록, 시장, 리서치, AI 리포트 화면을 제공합니다.",
            "Spring Boot 백엔드는 회원 인증, 거래 원장, 보유 종목, 시세 갱신, 리서치, 코치 메시지 API를 제공합니다.",
            "TransactionService가 거래 원장을 저장/수정/삭제할 때마다 포트폴리오를 재계산합니다.",
            "MarketDataProvider 인터페이스 뒤에 Yahoo provider와 Demo provider가 있어 외부 API 의존성을 낮춥니다.",
            "AI/코치 기능은 투자 조언이 아니라 리스크/체크포인트/학습 질문 생성에 초점을 둡니다.",
        ],
        "diagram": """
        [React Dashboard]
            | Bearer JWT / API calls
            v
        [Spring Boot API]
            | transaction create/update/delete
            v
        [TransactionService] -> rebuild holdings -> [PortfolioSnapshotService]
            |                                  |
            v                                  v
        [InvestmentTransaction]            [Holding / Portfolio]
            |
            v
        [MarketDataProvider: Yahoo or Demo] -> [Stock current price / PricePoint]
        """,
        "dirs": [
            ("backend", "Spring Boot API. auth, transaction, portfolio, marketdata, research, coach, ai 패키지로 나뉩니다."),
            ("frontend", "React 19 + TypeScript + Vite. API 모듈, 인증 컨텍스트, 페이지, 차트/초보자 컴포넌트가 있습니다."),
            ("docker-compose.yml", "DB 등 로컬 실행 환경을 구성하기 위한 파일입니다."),
            ("README.md", "서비스 목표, 주의사항, 실행 방법이 정리되어 있습니다."),
        ],
        "stack": [
            ("Frontend", "React 19, TypeScript, Vite 8, Tailwind 4, React Router, Axios, Recharts, lucide-react"),
            ("Backend", "Java 17, Spring Boot 4.0.6, Spring WebMVC, JPA, Security, Validation, Lombok"),
            ("Data", "H2 기본, MySQL profile, JPA entity/repository"),
            ("Market Data", "Provider abstraction, Yahoo Finance chart API, deterministic Demo fallback"),
            ("Auth", "직접 구현한 HS256 JWT provider와 Spring Security filter"),
        ],
        "domain": [
            ("InvestmentTransaction", "사용자 입력 거래 원장. BUY/SELL/DIVIDEND/DEPOSIT/WITHDRAWAL을 포함합니다."),
            ("Portfolio", "현금, 총자산, 평가금액, 총손익을 가진 집계 루트입니다."),
            ("Holding", "종목별 수량, 평균단가, 현재가, 평가금액, 비중을 가집니다."),
            ("Stock", "종목 마스터와 현재가, 섹터, 시가총액 정보를 가집니다."),
            ("MarketDataProvider", "시세 공급자를 교체할 수 있는 인터페이스입니다."),
            ("RiskAlert", "집중 비중, 손실/수익 종목 수, 최근 매수 증가 같은 점검 메시지입니다."),
        ],
        "flows": [
            {
                "title": "거래 기록 저장과 포트폴리오 재계산",
                "steps": [
                    "사용자가 거래 유형, 종목, 수량, 가격, 수수료, 세금, 이유, 태그를 입력합니다.",
                    "TransactionService.create가 현재 회원 ID와 종목을 확인하고 InvestmentTransaction을 저장합니다.",
                    "저장 후 rebuildPortfolioFromTransactions가 해당 회원의 보유 종목을 초기화합니다.",
                    "거래일/생성일 순서대로 원장을 다시 적용해 평균단가, 실현손익, 현금을 계산합니다.",
                    "PortfolioSnapshotService.refresh가 현재가 기준 평가금액과 비중을 갱신합니다.",
                ],
                "interview": "원장을 신뢰 가능한 단일 출처로 두고 파생 상태를 재계산하는 구조라고 설명하면 좋습니다.",
            },
            {
                "title": "시세 provider와 fallback",
                "steps": [
                    "MarketDataProvider 인터페이스가 providerName, supportsExternalApi, fetchQuotes를 정의합니다.",
                    "설정이 yahoo면 YahooFinanceMarketDataProvider가 chart API를 호출합니다.",
                    "외부 provider가 없으면 ConditionalOnMissingBean으로 DemoMarketDataProvider가 등록됩니다.",
                    "Demo provider는 시간/심볼 기반 deterministic drift로 화면과 계산을 유지합니다.",
                    "MarketDataRefreshService가 시세를 갱신하고 PricePoint를 저장합니다.",
                ],
                "interview": "외부 API가 없어도 개발/시연이 가능한 구조를 만든 것이 핵심입니다.",
            },
            {
                "title": "프론트 데이터 로딩과 mock fallback",
                "steps": [
                    "Axios 인스턴스가 access token을 Authorization 헤더에 붙입니다.",
                    "401 응답이면 localStorage token/member를 제거하고 auth-expired 이벤트를 발생시킵니다.",
                    "requestData는 API 실패 시 VITE_ALLOW_MOCK_FALLBACK 설정에 따라 mock 데이터를 반환합니다.",
                    "useAsyncData는 mock fallback이 꺼져 있으면 빈 런타임 값을 먼저 보여 주고 실패 이벤트를 발생시킵니다.",
                    "덕분에 실 API 연결 전에도 페이지 구조와 UX 검증이 가능합니다.",
                ],
                "interview": "시연 안정성을 위해 mock을 넣었지만, 운영에서는 fallback flag를 명시적으로 꺼야 한다고 말하면 좋습니다.",
            },
        ],
        "code": [
            {
                "file": "backend/src/main/java/com/stockflow/transaction/service/TransactionService.java",
                "snippet": """
                private void rebuildPortfolioFromTransactions(Long memberId) {
                    Portfolio portfolio = findPortfolio(memberId);
                    holdingRepository.deleteByPortfolioId(portfolio.getId());
                    portfolio.resetForRecordRecalculation(RECORD_BASE_CASH);

                    List<InvestmentTransaction> transactions = transactionRepository
                        .findByMemberIdOrderByTransactionDateDescCreatedAtDesc(memberId)
                        .stream()
                        .sorted(comparing(InvestmentTransaction::getTransactionDate)
                            .thenComparing(InvestmentTransaction::getCreatedAt))
                        .toList();

                    for (InvestmentTransaction tx : transactions) {
                        BigDecimal realized = applyRecord(portfolio, stock, tx.getTransactionType(), ...);
                        tx.updateCalculatedAmounts(totalAmount, realized);
                    }
                    portfolioSnapshotService.refresh(portfolio);
                }
                """,
                "explain": [
                    "거래 수정/삭제 후 보유 상태를 부분 보정하면 누락과 불일치가 생기기 쉽습니다.",
                    "이 코드는 거래 원장을 시간순으로 다시 적용해 파생 상태를 재생성합니다.",
                    "초기 기준 현금 1억 원을 두고 BUY는 출금, SELL은 입금/실현손익, DIVIDEND/DEPOSIT/WITHDRAWAL은 현금 변동으로 처리합니다.",
                    "면접에서는 이벤트 소싱의 간단한 형태라고 비유할 수 있습니다.",
                ],
            },
            {
                "file": "backend/src/main/java/com/stockflow/marketdata/provider/MarketDataProvider.java",
                "snippet": """
                public interface MarketDataProvider {
                    String providerName();
                    boolean supportsExternalApi();
                    List<StockQuoteSnapshot> fetchQuotes(List<Stock> stocks);
                }

                @ConditionalOnMissingBean(MarketDataProvider.class)
                public class DemoMarketDataProvider implements MarketDataProvider { ... }
                """,
                "explain": [
                    "시세 공급자를 인터페이스 뒤에 숨겨 Yahoo와 Demo를 교체 가능하게 했습니다.",
                    "ConditionalOnMissingBean은 외부 provider 설정이 없을 때 demo provider가 자동 등록되게 합니다.",
                    "개발, 테스트, 시연에서 외부 API 장애가 전체 서비스 실패로 번지지 않습니다.",
                    "새 provider를 붙일 때 controller/service 코드를 크게 바꾸지 않아도 됩니다.",
                ],
            },
            {
                "file": "frontend/src/hooks/useAsyncData.ts",
                "snippet": """
                export function useAsyncData<T>(loader, initialValue, deps = []) {
                  const mockFallbackEnabled = isMockFallbackEnabled();
                  const [data, setData] = useState(
                    mockFallbackEnabled ? initialValue : emptyRuntimeValue(initialValue)
                  );

                  loader()
                    .then(result => setData(result))
                    .catch(error => {
                      if (!mockFallbackEnabled) {
                        setData(emptyRuntimeValue(initialValue));
                        window.dispatchEvent(new CustomEvent("stockflow:data-error"));
                      }
                    });
                  return { data, setData, loading, error, placeholderData };
                }
                """,
                "explain": [
                    "API 연결 여부와 무관하게 같은 화면 컴포넌트를 재사용할 수 있게 만든 훅입니다.",
                    "fallback이 켜져 있으면 mock 데이터를 유지하고, 꺼져 있으면 빈 값과 에러 이벤트로 운영 상태를 드러냅니다.",
                    "데모 안정성과 실제 장애 가시성을 동시에 고려한 설계입니다.",
                    "다만 React Query/SWR 같은 검증된 캐시 레이어로 통합하는 것도 개선 방향입니다.",
                ],
            },
        ],
        "frontend": [
            "PortfolioPage는 portfolio, holdings, performance, allocation, transactions, studyCandidates를 각각 API에서 불러와 대시보드를 구성합니다.",
            "TransactionsPage는 wizard step, CSV preview/mapping, screenshot upload, memo, risk alert, filter를 한 화면에서 관리합니다.",
            "Axios request interceptor는 `stockflow.accessToken`을 붙이고, 401이면 저장된 세션 정보를 제거합니다.",
            "mock fallback은 시연에는 좋지만 실제 서비스에서는 환경변수로 명확히 끄는 운영 원칙이 필요합니다.",
        ],
        "security": [
            "JwtProvider는 HS256 서명, exp 검증, constant-time signature comparison을 직접 구현합니다.",
            "Spring Security는 `/api/auth/**`와 H2 console을 permitAll로 두고 나머지는 인증을 요구합니다.",
            "JWT filter는 Bearer token에서 memberId를 뽑아 SecurityContext principal로 넣습니다.",
            "투자 관련 기능은 법적 리스크가 있어 문서/리포트에 투자 조언 아님을 명확히 표시합니다.",
        ],
        "ops": [
            "MarketDataRefreshService는 provider status로 외부 API인지 demo인지 설명할 수 있게 합니다.",
            "Spring Boot build dir를 OS temp로 바꿔 OneDrive 동기화 환경에서 class 파일 문제가 덜 생기도록 했습니다.",
            "CSV import는 행 단위 오류를 모아 일부 성공/일부 실패를 사용자에게 설명합니다.",
        ],
        "tests": [
            "BUY/SELL/DIVIDEND/DEPOSIT/WITHDRAWAL 조합으로 포트폴리오 재계산 테스트가 필요합니다.",
            "같은 종목 다중 매수 후 일부 매도 시 평균단가와 실현손익을 검증해야 합니다.",
            "provider fallback, 401 인증 만료, mock fallback off 상태를 프론트/백엔드 테스트에 넣으면 좋습니다.",
        ],
        "strengths": [
            "투자 기록을 원장으로 보고 파생 상태를 재계산하는 구조가 명확합니다.",
            "시세 provider abstraction과 demo fallback 덕분에 외부 API 의존성이 낮습니다.",
            "투자 추천이 아니라 학습/체크리스트 도구라는 제품 경계가 비교적 분명합니다.",
        ],
        "improvements": [
            "직접 구현 JWT는 테스트가 충분해야 하며, 필요하면 검증된 라이브러리 사용도 고려할 수 있습니다.",
            "포트폴리오 재계산은 거래가 많아질수록 비용이 커지므로 snapshot/incremental update 전략을 고민해야 합니다.",
            "React Query로 API 캐시와 로딩/에러 상태를 일관화하면 프론트 코드가 더 단순해집니다.",
            "투자 리포트 문구는 법적 표현 검토와 disclaimer 위치를 더 엄격히 관리해야 합니다.",
        ],
        "qa": [
            ("왜 거래를 저장할 때마다 포트폴리오를 다시 계산하나요?", "거래 수정/삭제가 가능하기 때문에 파생 상태를 부분 수정하면 불일치 위험이 큽니다. 원장을 기준으로 재생성하면 정확성과 설명 가능성이 좋아집니다."),
            ("외부 시세 API가 실패하면 어떻게 하나요?", "Provider abstraction을 두고 외부 provider가 없거나 실패해도 demo provider/fallback으로 화면과 계산이 유지됩니다."),
            ("AI 기능은 투자 추천인가요?", "아닙니다. 기록을 바탕으로 리스크와 체크포인트를 정리하는 학습 보조 기능이며, 매수/매도 추천이나 투자 자문이 아니라고 명시했습니다."),
            ("확장 시 병목은?", "거래 수가 많을 때 전체 재계산 비용이 커질 수 있습니다. 이 경우 snapshot checkpoint나 event versioning을 도입할 수 있습니다."),
        ],
    },
    {
        "filename": "SignLanguage_면접대비_기술문서.docx",
        "title": "수어지구 Sign-Language",
        "subtitle": "한국어 문장을 수어 영상 재생 단위로 바꾸고, 수어 퀴즈/오답 학습을 제공하는 Expo + Spring Boot 서비스",
        "snapshot": repo_snapshot(
            "https://github.com/toadsam/Sign-Language",
            "5a9d4d7",
            "2026-05-28",
            "Use production Google OAuth client for Pages",
        ),
        "one_liner": (
            "수어지구는 한국어 문장을 수어 학습에 적합한 핵심 단어 순서로 단순화하고, 사전/Firebase Storage 영상과 연결해 "
            "학습자가 직접 수어 표현을 보고 퀴즈로 복습할 수 있게 만든 접근성 중심 서비스입니다."
        ),
        "problem": [
            "한국어 문장은 조사, 어미, 시제 표현이 많아 그대로 단어 검색을 하면 수어 영상 사전 매칭률이 낮아집니다.",
            "수어 학습자는 단어 영상뿐 아니라 문장 단순화 과정과 모르는 단어/영상 누락을 함께 알아야 합니다.",
            "외부 형태소 API와 OpenAI API는 실패할 수 있으므로 규칙 기반 fallback이 필요합니다.",
            "퀴즈는 Firebase/Firestore 데이터와 영상 Storage URL을 안정적으로 이어야 합니다.",
            "Expo 앱은 모바일 환경에서 OAuth, API base URL, 영상 재생 URL 보정이 중요합니다.",
        ],
        "architecture": [
            "Expo React Native 앱은 번역, 학습, 퀴즈, 오답 노트, Google auth 화면을 제공합니다.",
            "Spring Boot 백엔드는 `/translate`, `/api/quiz`, `/api/auth`, 사용자/오답 관련 API를 제공합니다.",
            "TranslationService는 규칙 기반 단순화, ETRI/외부 사전 후보, OpenAI 형태소 정규화를 조합합니다.",
            "DictionaryLoader는 `sign_dictionary.json`을 메모리에 올려 단어 -> 영상 메타데이터 매칭을 빠르게 합니다.",
            "StorageVideoCache는 Firebase Storage에서 단어 영상 URL을 찾아 캐시하고, 없으면 missing set에 기록합니다.",
        ],
        "diagram": """
        [Expo App]
           | POST /translate
           v
        [TranslationService]
           | rule simplifier
           | ETRI/KRDICT/Urimalsam candidates
           | OpenAI morphology normalization
           v
        [choose token stream by dictionary hits]
           | dictionary lookup + Firebase Storage URL cache
           v
        [playback items: word, file, url, hasVideo, unknown, noVideoWords]

        [Quiz Screen] -> /api/quiz/session -> [Firestore quiz_items + StorageVideoCache]
        """,
        "dirs": [
            ("backend", "Spring Boot API. translate, quiz, auth, dictionary, Firebase, external lexicon 서비스가 있습니다."),
            ("frontendcodes", "Expo/React Native 앱. app 라우트, context, lib/api, auth 모듈로 나뉩니다."),
            ("backend/src/main/resources", "application 설정과 sign_dictionary.json 같은 사전 리소스가 위치합니다."),
            ("frontendcodes/lib/api", "translate, quiz, auth, wrong notes 등 백엔드 호출 모듈입니다."),
        ],
        "stack": [
            ("Frontend", "Expo 54, React Native 0.81, expo-router, expo-auth-session, secure-store, video/image picker"),
            ("Backend", "Java 17, Spring Boot 3.2, Spring Security, OAuth2 client, Firebase Admin, Google API client"),
            ("AI/NLP", "OpenAI chat completions, ETRI WiseNLU, Urimalsam, KRDICT 후보 보강"),
            ("Storage", "Firebase Firestore quiz_items/users, Firebase Storage sign videos"),
            ("Auth", "Google ID token verification, JWT access/refresh generation"),
        ],
        "domain": [
            ("SignDictionaryEntry", "수어 사전의 단어, ID, 파일명을 나타내는 기본 매칭 단위입니다."),
            ("SimplificationResult", "단순화 문장, 토큰, 적용 규칙, 질문/부정/시제 메타데이터를 가집니다."),
            ("TranslatePlaybackItem", "단어별 영상 URL과 hasVideo 여부를 담아 앱 재생 순서를 구성합니다."),
            ("QuizItem", "Firestore quiz_items 문서. choices, correctChoiceId, videoUrl, level, stats를 포함합니다."),
            ("Wrong session", "사용자 문서의 incorrectQuestionCounts/incorrectQuestions를 기준으로 오답 퀴즈를 만듭니다."),
            ("GoogleUser", "Google ID token payload에서 뽑은 sub/email/name/picture 인증 주체입니다."),
        ],
        "flows": [
            {
                "title": "한국어 문장에서 수어 재생 목록까지",
                "steps": [
                    "앱이 `/translate`로 원문 텍스트를 보냅니다.",
                    "SignSentenceSimplifier가 조사 제거, 시제 분리, 의문사 재배치, 부정 처리, 단어 순서 재배열을 수행합니다.",
                    "ExternalLexiconApiClient가 ETRI 형태소 lemma와 우리말샘/KRDICT 후보를 가져옵니다.",
                    "OpenAiMorphologyNormalizerService가 가능하면 동사/형용사를 기본형으로 정규화한 JSON을 반환합니다.",
                    "TranslationService가 사전 매칭 수가 가장 높은 토큰 스트림을 선택하고 Firebase Storage 영상 URL을 붙입니다.",
                ],
                "interview": "AI가 실패해도 규칙/사전 기반으로 동작하고, AI 결과는 사전 매칭률로 검증해 선택한다는 점을 강조하세요.",
            },
            {
                "title": "퀴즈 세션과 정답 처리",
                "steps": [
                    "QuizService가 Firestore `quiz_items`에서 isActive=true 문서를 읽습니다.",
                    "카테고리 필터와 count 제한을 적용하고 문제를 섞습니다.",
                    "각 문제의 정답 choice text로 Firebase Storage 영상 URL을 찾고, 없으면 Firestore videoUrl fallback을 사용합니다.",
                    "정답 제출 시 correctChoiceId와 selectedChoiceId를 비교합니다.",
                    "attempt_count와 correct_count/wrong_count를 increment하되, 통계 갱신 실패는 정답 흐름을 막지 않습니다.",
                ],
                "interview": "학습 흐름은 항상 유지하고 통계 집계 실패는 부가 기능 실패로 격리한 설계입니다.",
            },
            {
                "title": "Google 인증과 토큰 발급",
                "steps": [
                    "Expo 앱이 Google OAuth로 ID token을 받습니다.",
                    "백엔드는 GoogleIdTokenVerifier로 audience와 서명을 검증합니다.",
                    "가입은 googleId 존재 여부를 확인해 신규 user를 만들고, 로그인은 기존 user를 조회합니다.",
                    "JwtService가 access token과 refresh token을 발급합니다.",
                    "JWT secret은 32자 이상이어야 하며 access/refresh 만료 기간은 설정으로 조정합니다.",
                ],
                "interview": "소셜 로그인은 클라이언트에서 받은 프로필을 믿는 것이 아니라 서버에서 ID token을 검증한다는 점이 중요합니다.",
            },
        ],
        "code": [
            {
                "file": "backend/src/main/java/com/wow/signlanguage/service/TranslationService.java",
                "snippet": """
                List<String> ruleTokens = simplification.tokens().stream()
                    .map(textNormalizer::normalizeToken)
                    .filter(token -> !token.isBlank())
                    .toList();

                List<String> etriTokens = externalLexiconApiClient.fetchSentenceLemmas(input);
                Optional<MorphologyNormalizationResult> openAiResult =
                    openAiMorphologyNormalizerService.normalize(input);

                TokenStreamChoice tokenChoice =
                    chooseTokenStream(ruleTokens, etriTokens, openAiTokens);
                List<String> resolvedTokens =
                    resolveTokens(tokenChoice.tokens(), mergeContextTokens(etriTokens, openAiTokens));
                """,
                "explain": [
                    "규칙 기반, 외부 형태소, OpenAI 결과를 모두 후보로 만든 뒤 사전 매칭률로 최종 토큰 스트림을 고릅니다.",
                    "AI 결과를 무조건 믿지 않고 실제 보유한 수어 사전에 얼마나 잘 맞는지를 기준으로 선택합니다.",
                    "unknown token은 외부 후보나 문맥 후보와 유사도 점수를 비교해 보정합니다.",
                    "결과에는 unknown과 noVideoWords를 나눠 사용자에게 '모르는 단어'와 '영상 누락'을 구분해 보여줄 수 있습니다.",
                ],
            },
            {
                "file": "backend/src/main/java/com/wow/signlanguage/service/SignSentenceSimplifier.java",
                "snippet": """
                orderedTokens.addAll(timeTokens);
                orderedTokens.addAll(placeTokens);
                orderedTokens.addAll(subjectTokens);
                orderedTokens.addAll(objectTokens);
                orderedTokens.addAll(predicateTokens);

                if (negative && !containsNegativeWord(orderedTokens)) {
                    addUnique(orderedTokens, "아니다");
                }
                if ("past".equals(tense) && hasDictionaryWord("끝")) {
                    addUnique(orderedTokens, "끝");
                }
                """,
                "explain": [
                    "수어 학습용 단순화는 조사/어미가 많은 한국어 문장을 의미 단위 중심으로 바꾸는 작업입니다.",
                    "시간, 장소, 주어, 목적어, 서술어 순서로 재배열해 영상 재생 순서를 안정화합니다.",
                    "부정과 과거 시제는 별도 토큰으로 분리해 사전 매칭 가능성을 높입니다.",
                    "규칙 기반이라 완벽하지 않지만 API 실패 시에도 예측 가능한 결과를 줍니다.",
                ],
            },
            {
                "file": "backend/src/main/java/com/wow/signlanguage/quiz/QuizService.java",
                "snippet": """
                String correctChoiceId = normalizeChoiceId(doc.getString("correctChoiceId"));
                List<String> choices = toStringList(doc.get("choices"));
                String correctChoiceText = choiceTextById(choices, correctChoiceId);
                String videoUrl = storageVideoCache.findUrlOrFallback(correctChoiceText, firestoreVideoUrl);

                updates.put("attempt_count", FieldValue.increment(1));
                updates.put(isCorrect ? "correct_count" : "wrong_count", FieldValue.increment(1));
                """,
                "explain": [
                    "퀴즈 영상은 정답 텍스트를 기준으로 Firebase Storage에서 다시 찾습니다.",
                    "Storage에 없으면 Firestore 문서의 videoUrl을 fallback으로 사용해 문제를 살릴 수 있습니다.",
                    "정답 통계는 FieldValue.increment로 원자적으로 증가시킵니다.",
                    "통계 갱신 실패를 무시하는 이유는 학습자의 정답 확인 흐름이 더 중요하기 때문입니다.",
                ],
            },
        ],
        "frontend": [
            "translate API는 clips와 items의 URL을 `resolveBackendUrl`로 보정해 로컬/배포 환경 차이를 줄입니다.",
            "quiz API는 session, wrong session, answer submit을 분리하고 영상 URL을 앱에서 재생 가능한 주소로 변환합니다.",
            "auth-context는 Google 로그인 결과와 secure store/token 흐름을 앱 전체에서 공유합니다.",
            "Expo 구조에서는 페이지 파일과 lib/api가 분리되어 화면과 네트워크 로직을 나누기 쉽습니다.",
        ],
        "security": [
            "Google ID token은 서버에서 audience 기준으로 검증합니다.",
            "JwtService는 secret 길이를 강제하고 access/refresh 만료를 분리합니다.",
            "현재 SecurityConfig는 `/api/quiz/**`, `/translate`, `/api/users/**`를 permitAll로 열어두므로 사용자별 오답/북마크 기능은 권한 범위 재점검이 필요합니다.",
            "외부 API key와 Firebase credential은 환경변수/secret으로 관리해야 합니다.",
        ],
        "ops": [
            "OpenAI 형태소 정규화는 10초 timeout과 Optional.empty fallback으로 전체 번역 실패를 막습니다.",
            "StorageVideoCache는 찾은 URL과 missing word를 캐시해 반복적인 Storage list 호출을 줄입니다.",
            "Firebase Storage 폴더 prefix 오타 후보를 둘 다 지원해 기존 데이터 구조와 호환합니다.",
        ],
        "tests": [
            "문장 단순화는 조사 제거, 의문문, 부정문, 과거/미래 시제별 예시를 고정 테스트해야 합니다.",
            "TranslationService는 rule/etri/openai 후보 중 사전 hit가 높은 쪽을 고르는 케이스를 테스트해야 합니다.",
            "QuizService는 Firestore 문서 필드 누락, choices 4개 미만, 영상 URL 없음, 통계 갱신 실패를 검증해야 합니다.",
        ],
        "strengths": [
            "AI를 보조 후보로 쓰고 실제 사전 매칭률로 결과를 검증하는 구조가 설득력 있습니다.",
            "unknown과 noVideoWords를 분리해 사용자에게 실패 원인을 구체적으로 보여줄 수 있습니다.",
            "퀴즈 통계 실패가 학습 흐름을 막지 않도록 부가 기능을 격리했습니다.",
        ],
        "improvements": [
            "permitAll 범위가 넓기 때문에 사용자별 데이터 API는 JWT 인증을 적용해야 합니다.",
            "OpenAI 응답 JSON 파싱 실패 케이스를 더 자세히 로깅하면 디버깅이 쉬워집니다.",
            "Storage list 조회는 데이터가 커질수록 비용이 커질 수 있으므로 단어->URL 인덱스를 별도 컬렉션으로 둘 수 있습니다.",
            "수어 문법은 규칙만으로 한계가 있어 실제 수어 전문가 피드백 기반 rule set 관리가 필요합니다.",
        ],
        "qa": [
            ("AI를 어디에 썼나요?", "한국어 문장의 형태소/기본형 정규화 후보를 얻는 데 사용했습니다. 다만 최종 결과는 사전 매칭률로 검증해 규칙/외부 API 결과와 비교합니다."),
            ("영상이 없는 단어는 어떻게 처리하나요?", "사전에 없는 단어는 unknown, 사전에는 있지만 Storage 영상이 없는 단어는 noVideoWords로 분리해 반환합니다."),
            ("퀴즈 통계 업데이트가 실패하면?", "정답 확인 응답은 계속 반환하고 통계 갱신 실패는 무시합니다. 학습 경험을 막지 않는 것이 우선입니다."),
            ("가장 큰 기술적 한계는?", "한국어 문장을 실제 수어 문법으로 완벽하게 바꾸는 것은 어려워 규칙/AI/사전의 하이브리드와 전문가 검수가 필요합니다."),
        ],
    },
    {
        "filename": "AjouCampusFoodmap_면접대비_기술문서.docx",
        "title": "Ajou Campus Foodmap",
        "subtitle": "아주대 주변 음식점 탐색, 사용자 제보, 관리자 운영을 제공하는 React + Express + MongoDB 서비스",
        "snapshot": (
            repo_snapshot("https://github.com/toadsam/pwd-week6-client", "2b5e44b", "2025-10-19", "3")
            + "\n"
            + repo_snapshot("https://github.com/toadsam/pwd-week6-server", "aba3b0a", "2026-05-07", "docs: add project README")
        ),
        "one_liner": (
            "Ajou Campus Foodmap은 아주대 주변 음식점을 목록/상세/인기 순으로 탐색하고, 로그인 사용자가 음식점을 제보하며, "
            "관리자가 음식점과 사용자 권한을 관리할 수 있게 만든 캠퍼스 맛집 서비스입니다."
        ),
        "problem": [
            "학생들은 학교 주변 음식점 정보를 흩어진 리뷰나 지도 앱에서 찾아야 합니다.",
            "서비스 운영자는 음식점 데이터 품질을 관리하고 사용자가 제보한 내용을 검토해야 합니다.",
            "로그인, 세션, 소셜 OAuth, 관리자 권한이 포함되어 단순 CRUD보다 인증 흐름이 중요합니다.",
            "서버는 MongoDB가 없을 때도 개발/테스트가 가능하도록 in-memory seed store를 지원합니다.",
            "클라이언트는 공개 조회와 세션 기반 인증 요청을 분리해야 CORS/cookie 문제가 줄어듭니다.",
        ],
        "architecture": [
            "React 클라이언트는 음식점 목록/상세/인기, 제출, 로그인/회원가입, 관리자 페이지를 제공합니다.",
            "Express 서버는 restaurant, submission, auth, users 라우터와 controller/service/model 계층으로 구성됩니다.",
            "MongoDB 연결 시 Mongoose 모델을 사용하고, 연결이 없으면 restaurant seed JSON을 in-memory store로 사용합니다.",
            "Passport Local/Google/Naver 전략과 express-session으로 세션 기반 인증을 처리합니다.",
            "React Query는 관리자 페이지에서 음식점/사용자 데이터를 캐시하고 mutation 후 invalidate합니다.",
        ],
        "diagram": """
        [React Client]
          | publicApi without credentials
          v
        GET /api/restaurants -> [Express routes/controllers] -> [restaurants.service]
                                                         | DB connected?
                                                         +--> [MongoDB/Mongoose]
                                                         +--> [in-memory seed store]

        [React Login/Admin]
          | api withCredentials
          v
        /api/auth + /api/users + /api/submissions -> [Passport session + auth middleware]
        """,
        "dirs": [
            ("Foodmap-client/src/services", "axios API wrapper와 auth API 모듈이 있습니다."),
            ("Foodmap-client/src/contexts", "AuthContext가 현재 사용자와 로그인 상태를 관리합니다."),
            ("Foodmap-client/src/pages", "List, Detail, Submit, Admin 같은 화면 단위가 있습니다."),
            ("Foodmap-server/src/routes", "auth, users, restaurants, submissions 라우터를 정의합니다."),
            ("Foodmap-server/src/services", "DB/in-memory 접근과 비즈니스 로직을 controller에서 분리합니다."),
            ("Foodmap-server/tests", "restaurants service/routes Jest + Supertest 테스트가 포함됩니다."),
        ],
        "stack": [
            ("Frontend", "React 19, Vite 7, Emotion, React Query 5, Axios, React Hook Form, React Router 7"),
            ("Backend", "Node.js, Express 5, Mongoose 8, express-session, connect-mongo"),
            ("Auth", "Passport Local, Google OAuth, Naver OAuth, bcryptjs password hashing"),
            ("Data", "MongoDB, Restaurant/Submission/User Mongoose models, JSON seed fallback"),
            ("Test", "Jest, Supertest, mongodb-memory-server"),
        ],
        "domain": [
            ("Restaurant", "id, name, category, location, priceRange, rating, description, recommendedMenu, likes, image를 가집니다."),
            ("Submission", "사용자 제보 데이터. pending 상태로 생성되고 관리자가 조회/수정/삭제합니다."),
            ("User", "local/google/naver provider, providerId, email, name, avatar, userType, isActive를 가집니다."),
            ("Session", "connect.sid cookie와 Passport serialize/deserializeUser로 유지됩니다."),
            ("Admin", "userType이 admin인 사용자로 submissions/users 관리 권한을 가집니다."),
        ],
        "flows": [
            {
                "title": "음식점 목록 조회",
                "steps": [
                    "클라이언트 `restaurantAPI.getRestaurants`가 publicApi로 `/api/restaurants`를 호출합니다.",
                    "publicApi는 withCredentials=false라 공개 조회에서 쿠키/CORS 문제를 줄입니다.",
                    "Express router가 restaurantsController.getRestaurants로 요청을 넘깁니다.",
                    "restaurants.service는 DB 연결 상태를 확인해 MongoDB 또는 in-memory store에서 데이터를 가져옵니다.",
                    "컨트롤러는 `{ data: restaurants }` 형태로 응답하고 화면은 목록/카드로 렌더링합니다.",
                ],
                "interview": "DB가 없어도 seed 기반 개발이 가능하게 서비스 계층에서 저장소 선택을 숨긴 점을 설명하세요.",
            },
            {
                "title": "세션 기반 로그인과 현재 사용자 확인",
                "steps": [
                    "사용자가 이메일/비밀번호로 `/api/auth/login`을 호출합니다.",
                    "Passport local strategy가 provider=local 사용자와 bcrypt password를 확인합니다.",
                    "req.login이 성공하면 express-session이 connect.sid 쿠키를 내려줍니다.",
                    "프론트 AuthContext는 마운트 시 `/api/auth/me`를 호출해 세션을 복원합니다.",
                    "로그아웃은 req.logout, req.session.destroy, clearCookie 순서로 세션을 정리합니다.",
                ],
                "interview": "JWT가 아니라 서버 세션을 쓴 이유와, CORS + secure cookie + sameSite 설정을 함께 설명하면 좋습니다.",
            },
            {
                "title": "사용자 제보와 관리자 처리",
                "steps": [
                    "로그인 사용자가 제보 폼으로 restaurantName, category, location, menu, review를 보냅니다.",
                    "submissions route는 create에 isAuthenticated를 적용해 로그인 사용자만 제보하게 합니다.",
                    "controller는 필수 필드를 검증하고 status를 pending으로 고정합니다.",
                    "관리자 route는 isAdmin을 통해 list/get/update/delete를 제한합니다.",
                    "AdminPage는 React Query mutation 후 invalidateQueries로 최신 목록을 다시 불러옵니다.",
                ],
                "interview": "사용자 생성 데이터와 운영자 승인 데이터를 분리한 구조라고 말하면 좋습니다.",
            },
        ],
        "code": [
            {
                "file": "Foodmap-server/src/services/restaurants.service.js",
                "snippet": """
                function usingDb() {
                  return mongoose.connection && mongoose.connection.readyState === 1;
                }

                async function getAllRestaurants() {
                  if (usingDb()) {
                    const docs = await Restaurant.find().select("-_id").lean();
                    cachedAll = docs;
                    return deepClone(docs);
                  }
                  return deepClone(memStore);
                }
                """,
                "explain": [
                    "서비스 계층이 DB 연결 여부를 판단해 MongoDB와 in-memory store를 전환합니다.",
                    "테스트/개발 환경에서 MongoDB가 없어도 기본 기능을 확인할 수 있습니다.",
                    "deepClone으로 외부에서 내부 store를 직접 변형하지 못하게 합니다.",
                    "운영에서는 DB 연결 실패를 조용히 fallback할지, 명확히 장애로 볼지 정책을 정해야 합니다.",
                ],
            },
            {
                "file": "Foodmap-server/src/config/passport.config.js",
                "snippet": """
                passport.serializeUser((user, done) => done(null, user._id));
                passport.deserializeUser(async (id, done) => {
                  const user = await User.findById(id);
                  done(null, user);
                });

                passport.use(new LocalStrategy({ usernameField: "email" }, async (email, password, done) => {
                  const user = await User.findOne({ email: email.toLowerCase(), provider: "local" });
                  if (!user || !(await user.comparePassword(password))) return done(null, false);
                  return done(null, user);
                }));
                """,
                "explain": [
                    "Passport는 세션에 사용자 전체가 아니라 user ID만 저장합니다.",
                    "요청마다 deserializeUser가 DB에서 사용자 객체를 복원합니다.",
                    "local strategy는 provider=local 조건을 걸어 OAuth 사용자와 이메일 충돌을 줄입니다.",
                    "Google/Naver OAuth도 provider/providerId를 기준으로 기존 사용자 조회 후 없으면 생성합니다.",
                ],
            },
            {
                "file": "Foodmap-server/src/routes/restaurants.routes.js",
                "snippet": """
                router.get("/popular", restaurantsController.getPopularRestaurants);
                router.get("/", restaurantsController.getRestaurants);
                router.get("/:id", restaurantsController.getRestaurant);

                // 관리자 전용 엔드포인트
                router.post("/", restaurantsController.createRestaurant);
                router.put("/:id", restaurantsController.updateRestaurant);
                router.delete("/:id", restaurantsController.deleteRestaurant);
                """,
                "explain": [
                    "주석은 관리자 전용이라고 되어 있지만 실제로 isAdmin 미들웨어가 붙어 있지 않습니다.",
                    "반면 submissions route는 create에 isAuthenticated, 관리 기능에 isAdmin이 적용되어 있습니다.",
                    "면접에서는 이 부분을 숨기기보다 '리뷰 중 발견한 개선점'으로 말하는 편이 더 신뢰도 높습니다.",
                    "수정 방향은 `const { isAdmin } = require('../middleware/auth.middleware')` 후 post/put/delete에 적용하는 것입니다.",
                ],
            },
        ],
        "frontend": [
            "api.jsx는 인증 요청용 `api(withCredentials=true)`와 공개 조회용 `publicApi(withCredentials=false)`를 분리합니다.",
            "AuthContext는 `/api/auth/me`로 자동 로그인 상태를 확인하고 user/isAuthenticated/isLoading을 제공합니다.",
            "AdminPage는 React Query로 restaurants/users를 불러오고, mutation 성공 시 invalidateQueries로 목록을 갱신합니다.",
            "React Hook Form은 관리자 음식점 생성/수정 폼의 입력 상태를 관리합니다.",
        ],
        "security": [
            "세션 쿠키는 httpOnly, production secure, sameSite none/lax 설정을 환경에 따라 바꿉니다.",
            "프록시 환경에서 secure cookie가 정상 설정되도록 `app.set('trust proxy', 1)`을 사용합니다.",
            "submissions/users에는 인증/관리자 미들웨어가 있지만 restaurants CUD에는 누락되어 있습니다.",
            "OAuth providerId unique 충돌을 피하기 위해 local 계정도 providerId를 이메일로 설정합니다.",
        ],
        "ops": [
            "MongoDB 연결이 있으면 MongoStore로 세션을 저장하고, 없으면 메모리 세션을 사용합니다.",
            "health endpoint는 DB 연결 상태, timestamp, environment를 반환합니다.",
            "seed JSON을 통해 초기 음식점 데이터를 빠르게 복원할 수 있습니다.",
        ],
        "tests": [
            "restaurants service/routes 테스트가 있어 목록/상세/생성/수정/삭제 기본 회귀를 잡을 수 있습니다.",
            "auth middleware는 로그인 안 된 사용자, 일반 사용자, admin 사용자 케이스를 테스트해야 합니다.",
            "restaurants CUD에 isAdmin을 붙인 뒤 Supertest로 401/403/성공 케이스를 추가하는 것이 좋습니다.",
        ],
        "strengths": [
            "controller/service/model 계층이 분리되어 Express 프로젝트 구조를 설명하기 좋습니다.",
            "MongoDB와 in-memory seed fallback이 있어 개발/테스트 진입 장벽이 낮습니다.",
            "세션 기반 인증과 OAuth provider 처리 경험을 보여줄 수 있습니다.",
        ],
        "improvements": [
            "restaurant 생성/수정/삭제 route에 isAdmin 미들웨어를 반드시 붙여야 합니다.",
            "메모리 세션은 운영에 부적합하므로 production에서는 MongoStore 필수 검증을 넣는 편이 좋습니다.",
            "API 응답 형식이 `{data}`와 `{success,message,data}`로 섞여 있어 표준화하면 프론트 코드가 단순해집니다.",
            "client repo에 node_modules가 포함된 흔적은 정리하고 package-lock 기반 복원을 권장합니다.",
        ],
        "qa": [
            ("JWT가 아니라 세션을 쓴 이유는?", "Passport와 OAuth 흐름을 빠르게 붙이고, 브라우저 쿠키 기반으로 로그인 상태를 유지하기 위해 세션 방식을 사용했습니다."),
            ("DB가 없으면 어떻게 동작하나요?", "restaurants.service가 MongoDB 연결 상태를 확인하고, 연결이 없으면 JSON seed를 읽어 in-memory store로 동작합니다."),
            ("가장 중요한 보안 개선점은?", "restaurant POST/PUT/DELETE에 관리자 미들웨어가 빠져 있어 `isAdmin`을 붙이고 테스트를 추가해야 합니다."),
            ("React Query를 왜 썼나요?", "관리자 페이지에서 생성/수정/삭제 후 invalidateQueries로 서버 상태를 다시 동기화하기 쉽기 때문입니다."),
        ],
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run, size=9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, font: str = "Malgun Gothic") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_code(doc: Document, code: str) -> None:
    cleaned = dedent(code).strip("\n")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.05)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(cleaned)
    set_run_font(run, size=8.2, font="Consolas")
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F6FA")
    p_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_section(doc: Document, title: str, lead: str | None = None) -> None:
    doc.add_page_break()
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        set_run_font(run, size=16, bold=True, color="2E74B5")
    if lead:
        add_paragraph(doc, lead)


def add_subheading(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        set_run_font(run, size=13, bold=True, color="2E74B5")


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for list_style in ["List Bullet", "List Number"]:
        style = doc.styles[list_style]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)

    return doc


def add_cover(doc: Document, project: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run(project["title"])
    set_run_font(run, size=24, bold=True, color="1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(project["subtitle"])
    set_run_font(run, size=12.5, color="333333")

    add_table(
        doc,
        ["항목", "내용"],
        [
            ("문서 목적", "면접 대비용 기술 설명서. 구조, 핵심 코드, 동작 원리, 예상 질문 답변을 한 문서로 정리합니다."),
            ("분량 의도", "각 장을 독립 페이지처럼 읽을 수 있게 구성해 20페이지 내외의 발표/복습용 문서로 만들었습니다."),
            ("소스 스냅샷", project["snapshot"]),
            ("추천 사용법", "1분 소개 -> 구조 설명 -> 핵심 코드 3개 -> 한계/개선 -> 예상 질문 순서로 연습합니다."),
        ],
    )
    add_paragraph(
        doc,
        "면접에서는 모든 기능을 나열하기보다, 문제 정의와 설계 판단을 먼저 말하고 그 뒤에 코드 근거를 붙이는 순서가 좋습니다. "
        "이 문서는 그 흐름으로 바로 말할 수 있게 구성했습니다.",
    )


def add_intro(doc: Document, project: dict) -> None:
    add_section(doc, "1. 프로젝트 소개와 발표 방향", project["one_liner"])
    add_subheading(doc, "소개할 때 잡을 중심축")
    add_bullets(doc, [
        "서비스가 해결하려는 문제를 먼저 말합니다.",
        "사용자 경험과 기술 구조가 어떻게 연결되는지 설명합니다.",
        "핵심 구현 3개를 코드 파일명과 함께 말할 준비를 합니다.",
        "한계와 개선점을 숨기지 않고 운영 관점으로 정리합니다.",
    ])
    add_subheading(doc, "30초 버전")
    add_paragraph(doc, project["one_liner"])
    add_paragraph(
        doc,
        "제가 이 프로젝트를 설명한다면 기능 목록보다 '왜 이런 구조가 필요했는가'를 먼저 잡겠습니다. "
        "그 다음 API 흐름, 상태 관리, 인증/실시간/외부 연동 같은 어려운 부분을 실제 코드와 연결해 설명합니다.",
    )


def add_problem(doc: Document, project: dict) -> None:
    add_section(doc, "2. 문제 정의와 사용자 시나리오")
    add_bullets(doc, project["problem"])
    add_subheading(doc, "면접 답변 포인트")
    add_paragraph(
        doc,
        "문제 정의는 프로젝트의 설계 선택을 납득시키는 출발점입니다. 단순히 '무엇을 만들었다'가 아니라 "
        "'어떤 사용자가 어떤 상황에서 불편했고, 그래서 어떤 데이터와 흐름이 필요했다'라고 말해야 합니다.",
    )
    add_paragraph(
        doc,
        "이 프로젝트를 소개할 때는 사용자 액션 하나가 백엔드, 데이터, 화면 상태, 예외 처리로 어떻게 이어지는지 "
        "하나의 시나리오로 묶어 말하는 것이 좋습니다.",
    )


def add_architecture(doc: Document, project: dict) -> None:
    add_section(doc, "3. 전체 아키텍처 한 장 요약")
    add_bullets(doc, project["architecture"])
    add_subheading(doc, "흐름도")
    add_code(doc, project["diagram"])
    add_subheading(doc, "설명 순서")
    add_numbered(doc, [
        "클라이언트가 어떤 사용자 액션을 받는지 설명합니다.",
        "API 또는 실시간 채널로 어떤 데이터가 이동하는지 설명합니다.",
        "백엔드 서비스가 어떤 도메인 객체를 변경하거나 조회하는지 설명합니다.",
        "실패했을 때 fallback 또는 오류 처리가 어디에 있는지 설명합니다.",
    ])


def add_structure(doc: Document, project: dict) -> None:
    add_section(doc, "4. 저장소 구조와 읽는 순서")
    add_table(doc, ["폴더/파일", "역할"], project["dirs"])
    add_subheading(doc, "코드를 처음 읽는 순서")
    add_numbered(doc, [
        "README와 빌드 파일을 먼저 보고 실행 단위와 의존성을 파악합니다.",
        "라우터/컨트롤러에서 외부 API 표면을 확인합니다.",
        "서비스 계층에서 핵심 비즈니스 규칙과 트랜잭션 경계를 읽습니다.",
        "프론트 API 모듈과 화면 컴포넌트가 응답을 어떻게 소비하는지 확인합니다.",
        "마지막으로 인증, 예외, 외부 연동, 테스트를 보며 운영 리스크를 찾습니다.",
    ])


def add_stack(doc: Document, project: dict) -> None:
    add_section(doc, "5. 기술 스택과 선택 이유")
    add_table(doc, ["영역", "사용 기술 / 의미"], project["stack"])
    add_subheading(doc, "면접에서 말할 방식")
    add_paragraph(
        doc,
        "기술 스택은 이름을 나열하는 것보다 '이 기술이 이 프로젝트의 어떤 문제를 해결했는가'로 설명해야 합니다. "
        "예를 들어 인증은 보안 상태 유지, SSE/Socket은 실시간 UX, provider abstraction은 외부 API 실패 대응처럼 연결합니다.",
    )
    add_paragraph(
        doc,
        "모르는 세부 버전 질문을 받으면 버전 번호를 외우는 것보다 주요 API와 설계 이유를 중심으로 답변하는 편이 안정적입니다.",
    )


def add_domain(doc: Document, project: dict) -> None:
    add_section(doc, "6. 도메인 모델과 책임 분리")
    add_table(doc, ["도메인/개념", "역할"], project["domain"])
    add_subheading(doc, "책임 분리 원칙")
    add_bullets(doc, [
        "컨트롤러는 요청/응답 변환과 상태 코드에 집중합니다.",
        "서비스는 비즈니스 규칙과 여러 저장소 호출의 순서를 책임집니다.",
        "모델/엔티티는 데이터 형태와 최소한의 상태 변경 메서드를 가집니다.",
        "프론트 API 모듈은 HTTP 세부 사항을 감추고 화면은 사용자 흐름에 집중합니다.",
    ])


def add_flow(doc: Document, idx: int, flow: dict) -> None:
    add_section(doc, f"{6 + idx}. 핵심 흐름 {idx}: {flow['title']}")
    add_numbered(doc, flow["steps"])
    add_subheading(doc, "원리")
    add_paragraph(
        doc,
        "이 흐름은 단일 함수 하나보다 여러 레이어의 협업으로 이해해야 합니다. 면접에서는 입력, 검증, 상태 변경, "
        "응답 또는 브로드캐스트, 실패 처리 순서로 말하면 구조가 분명해집니다.",
    )
    add_subheading(doc, "면접에서 이렇게 말하기")
    add_paragraph(doc, flow["interview"])


def add_code_section(doc: Document, idx: int, item: dict) -> None:
    add_section(doc, f"{9 + idx}. 중요 코드 {idx}: {item['file']}")
    add_code(doc, item["snippet"])
    add_subheading(doc, "코드 해설")
    add_bullets(doc, item["explain"])
    add_subheading(doc, "꼬리 질문 대비")
    add_paragraph(
        doc,
        "이 코드를 설명할 때는 문법 설명에 오래 머무르지 말고, 왜 이 위치에 이 로직이 있어야 하는지 말해야 합니다. "
        "그리고 같은 문제를 다른 방식으로 풀 수 있는 대안과 현재 방식의 한계를 같이 말하면 답변의 신뢰도가 올라갑니다.",
    )


def add_frontend(doc: Document, project: dict) -> None:
    add_section(doc, "13. 프론트엔드 구조와 상태 관리")
    add_bullets(doc, project["frontend"])
    add_subheading(doc, "화면 설명 순서")
    add_numbered(doc, [
        "사용자 화면에서 어떤 행동을 하는지 말합니다.",
        "해당 행동이 어떤 API 모듈 또는 실시간 채널을 호출하는지 연결합니다.",
        "응답 데이터가 어떤 컴포넌트 상태나 캐시에 반영되는지 설명합니다.",
        "로딩, 실패, 권한 만료, fallback 화면을 어떻게 처리하는지 덧붙입니다.",
    ])


def add_security(doc: Document, project: dict) -> None:
    add_section(doc, "14. 인증, 보안, 오류 처리")
    add_bullets(doc, project["security"])
    add_subheading(doc, "면접 답변 태도")
    add_paragraph(
        doc,
        "보안은 '완벽하다'고 말하기보다 현재 적용된 방어와 남은 위험을 나눠 말해야 합니다. "
        "토큰, 세션, 권한, CORS, 외부 key, 민감 API 접근 제어를 체크리스트처럼 짚으면 좋습니다.",
    )


def add_ops(doc: Document, project: dict) -> None:
    add_section(doc, "15. 데이터, 외부 연동, 운영 관점")
    add_bullets(doc, project["ops"])
    add_subheading(doc, "운영 관점에서 볼 질문")
    add_bullets(doc, [
        "외부 API가 느리거나 실패하면 사용자는 어떤 화면을 보게 되는가?",
        "캐시가 오래되거나 누락되면 어떤 데이터가 틀릴 수 있는가?",
        "실시간 연결 또는 세션이 끊기면 자동 복구되는가?",
        "로그만 보고 장애 원인을 추적할 수 있는가?",
    ])


def add_tests(doc: Document, project: dict) -> None:
    add_section(doc, "16. 테스트와 검증 전략")
    add_bullets(doc, project["tests"])
    add_subheading(doc, "우선순위")
    add_paragraph(
        doc,
        "테스트는 모든 줄을 덮는 것보다 비즈니스 결과가 틀리면 위험한 흐름을 먼저 잡아야 합니다. "
        "인증, 권한, 계산, 외부 연동 fallback, 날짜/시간 경계, 동시성은 면접에서도 자주 나오는 주제입니다.",
    )


def add_strengths(doc: Document, project: dict) -> None:
    add_section(doc, "17. 강점, 한계, 개선 답변")
    add_subheading(doc, "강점")
    add_bullets(doc, project["strengths"])
    add_subheading(doc, "한계와 개선")
    add_bullets(doc, project["improvements"])
    add_subheading(doc, "말하는 방식")
    add_paragraph(
        doc,
        "한계는 약점이 아니라 설계 판단의 경계를 보여주는 재료입니다. '현재는 MVP라 이렇게 했고, 운영 단계에서는 "
        "이 순서로 보완하겠다'라고 답하면 방어적인 답변보다 훨씬 좋습니다.",
    )


def add_qa(doc: Document, project: dict) -> None:
    add_section(doc, "18. 면접 예상 질문과 답변")
    for q, a in project["qa"]:
        add_subheading(doc, f"Q. {q}")
        add_paragraph(doc, f"A. {a}")
    add_subheading(doc, "공통 꼬리 질문")
    add_bullets(doc, [
        "이 기능을 처음부터 다시 만든다면 무엇을 바꾸겠나요?",
        "장애가 난다면 어디부터 로그를 보겠나요?",
        "데이터가 10배 늘어나면 어떤 부분이 먼저 병목이 될까요?",
        "본인이 직접 구현한 부분과 참고한 부분을 어떻게 구분해서 설명할 수 있나요?",
    ])


def add_script(doc: Document, project: dict) -> None:
    add_section(doc, "19. 발표 스크립트와 최종 체크리스트")
    add_subheading(doc, "1분 소개")
    add_paragraph(
        doc,
        f"{project['title']}은 {project['one_liner']} "
        "제가 소개할 때는 사용자 문제, 전체 구조, 핵심 코드 3개, 한계와 개선 순서로 설명하겠습니다.",
    )
    add_subheading(doc, "3분 소개")
    add_paragraph(
        doc,
        "먼저 문제 상황을 말합니다. 사용자는 단순 조회나 기록을 넘어 지속적인 피드백과 안정적인 사용 경험을 기대합니다. "
        "그래서 프론트엔드는 화면 흐름과 상태 관리를 맡고, 백엔드는 인증과 도메인 규칙, 데이터 저장, 외부 연동을 맡도록 나눴습니다. "
        "핵심 구현은 첫째, 사용자 액션이 서비스 계층에서 어떤 순서로 처리되는지, 둘째, 인증과 권한이 어디서 검증되는지, "
        "셋째, 외부 API나 실시간 연결 실패를 어떻게 다루는지입니다. 마지막으로 현재 구조의 한계와 운영 단계 개선 방향까지 말하겠습니다.",
    )
    add_subheading(doc, "5분 소개")
    add_paragraph(
        doc,
        "5분 발표에서는 코드 파일명을 직접 언급하면서 말합니다. 먼저 저장소 구조를 짧게 보여주고, 핵심 flow 하나를 end-to-end로 설명합니다. "
        "그 뒤 중요 코드 1은 인증/권한, 중요 코드 2는 핵심 도메인 계산 또는 실시간 처리, 중요 코드 3은 fallback/상태 관리로 잡습니다. "
        "각 코드는 '무엇을 하는가'보다 '왜 이 방식이 필요한가'를 중심으로 설명합니다. 마지막 30초는 개선점입니다. "
        "운영 전 보완할 보안, 테스트, 확장성 항목을 말하면 프로젝트를 객관적으로 이해하고 있다는 인상을 줄 수 있습니다.",
    )
    add_subheading(doc, "최종 체크리스트")
    add_bullets(doc, [
        "프로젝트 한 줄 정의를 외운다.",
        "사용자 문제 2개와 기술 문제 2개를 구분해서 말한다.",
        "핵심 파일 3개를 경로와 함께 말할 수 있다.",
        "데이터가 입력되어 화면에 반영되기까지의 흐름을 설명할 수 있다.",
        "보안/권한/실패 처리의 현재 상태와 개선점을 말할 수 있다.",
        "내가 직접 구현한 부분과 팀/라이브러리/외부 API의 역할을 구분한다.",
    ])


def build_doc(project: dict) -> Path:
    doc = setup_document()
    add_cover(doc, project)
    add_intro(doc, project)
    add_problem(doc, project)
    add_architecture(doc, project)
    add_structure(doc, project)
    add_stack(doc, project)
    add_domain(doc, project)
    for idx, flow in enumerate(project["flows"], start=1):
        add_flow(doc, idx, flow)
    for idx, item in enumerate(project["code"], start=1):
        add_code_section(doc, idx, item)
    add_frontend(doc, project)
    add_security(doc, project)
    add_ops(doc, project)
    add_tests(doc, project)
    add_strengths(doc, project)
    add_qa(doc, project)
    add_script(doc, project)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUT_DIR / project["filename"]
    doc.save(output)
    return output


def main() -> None:
    outputs = []
    for project in PROJECTS:
        outputs.append(build_doc(project))
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
